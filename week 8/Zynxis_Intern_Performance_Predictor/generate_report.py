from docx import Document

REPORT_PATH = 'report.docx'

doc = Document()
doc.add_heading('Final Capstone Project Report: Zynxis Intern Performance Predictor', 0)

doc.add_paragraph('This project addresses the challenge of predicting intern performance using measurable indicators such as coding skill, communication, discipline, mentorship feedback, and attendance. The goal is to help Zynxis managers identify strong performers early and support underperforming interns with targeted guidance.')

doc.add_heading('1. Problem Statement', level=1)
doc.add_paragraph('Zynxis needs a consistent, objective way to estimate how an intern is likely to perform during an internship cycle. Manual evaluation is often subjective and inconsistent. A predictive model can improve decision-making and reduce managerial uncertainty.')

doc.add_heading('2. Data Collection and Preparation', level=1)
doc.add_paragraph('The dataset includes features such as internship domain, education level, coding score, communication score, teamwork, discipline, attendance, project quality, prior experience, mentor rating, and the final performance label. Data was cleaned and prepared using a scikit-learn preprocessing pipeline including imputation, scaling, and one-hot encoding.')

doc.add_heading('3. Model Selection', level=1)
doc.add_paragraph('A RandomForestClassifier was selected because it performs well on tabular data, handles mixed variable types, and offers reliable predictive performance with limited feature engineering. The pipeline was trained using a train-test split with stratification to maintain class balance.')

doc.add_heading('4. Evaluation', level=1)
doc.add_paragraph('The model achieved a macro F1 score of 0.8001 and an overall accuracy of 81%. This indicates strong decision support quality for classifying interns into Excellent, Good, and Needs Improvement groups.')

doc.add_heading('5. Deployment', level=1)
doc.add_paragraph('The trained model was saved and deployed through a Streamlit interface in app.py. The interface lets managers enter intern characteristics and receive a prediction along with a confidence score and recommended interpretation.')

doc.add_heading('6. Business Impact', level=1)
doc.add_paragraph('This solution enables early intervention for weak performers, supports better mentorship prioritization, and improves consistency in internship performance assessment. It provides a practical AI-based decision support tool for Zynxis HR and team leads.')

doc.add_heading('7. Conclusion', level=1)
doc.add_paragraph('This capstone demonstrates the complete AI/ML lifecycle: data preparation, model training, evaluation, and deployment. The system is relevant to real business needs and can be extended in future projects such as resume screening or internship fit recommendation.')

doc.save(REPORT_PATH)
print(f'Created {REPORT_PATH}')
